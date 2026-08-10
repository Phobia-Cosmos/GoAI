from __future__ import annotations

import json
import os
import re
import sqlite3
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
import xlrd
from docx import Document
from openpyxl import load_workbook

from .common import (
    as_int,
    as_number,
    cell_ref,
    clean_text,
    is_blank,
    parse_duration,
    parse_money_wan,
    parse_period,
    sha256_file,
    source_id,
    stable_record_id,
)
from .rulepack import (
    ACTION_CONTROL_TYPES,
    ACTION_DEFINITIONS,
    AGENT_CANDIDATE_ACTIONS,
    HISTORICAL_ACTION_ALIASES,
    RULE_GAPS,
    TEST_ACTION_ALIASES,
    canonicalize_action,
    parse_action_parameters,
)


HISTORICAL_COMPETITION = "historical_600_unknown_rule"
TEST_COMPETITION = "zhejiang_8th_710"
ORDER_COMPETITION = "order_catalog_unbound"
DESIGN_SCOPE = "design_reference"

MARKET_NAMES = {1: "本地", 2: "区域", 3: "国内", 4: "亚洲", 5: "国际"}
PRODUCT_NAMES = {index: f"P{index}" for index in range(1, 10)}
ISO_NAMES = {0: None, 1: "ISO9000", 2: "ISO14000"}


TABLE_DESCRIPTIONS = {
    "source_manifest": "源文件清单、哈希、批次归属和重复关系",
    "workbook_inventory": "所有 Excel 工作簿及工作表结构",
    "workbook_cells": "所有 Excel 非空单元格的保真长表",
    "competition_rules": "题面级比赛元数据",
    "rule_financing": "融资与贴现规则",
    "rule_financing_terms": "可审计的融资与贴现分档条款",
    "rule_factories": "厂房规则",
    "rule_production_lines": "生产线规则",
    "rule_markets": "市场开拓规则",
    "rule_iso": "ISO 认证规则",
    "rule_materials": "原材料规则",
    "rule_products": "产品研发及成本规则",
    "rule_bom": "产品物料清单",
    "rule_packs": "规则包版本、绑定和可执行状态",
    "rule_gaps": "阻止完整仿真或需要确认的规则缺口",
    "action_definitions": "跨来源标准动作定义与观察覆盖",
    "action_aliases": "原始动作名称到标准动作的映射",
    "action_events": "历史与测试现金流统一后的结构化经营事件",
    "order_catalog": "赛前订单目录",
    "annual_advertising": "公共年度广告巡盘长表",
    "annual_financial_metrics": "公共年度三张报表长表",
    "annual_production_lines": "公共年度产线巡盘",
    "annual_market_leaders": "公共年度市场老大",
    "teams": "企业基本信息与数据可用状态",
    "team_cash_flows": "企业现金流水明细",
    "team_financial_metrics": "企业历年三张报表长表",
    "team_advertising": "企业历年广告投放长表",
    "team_orders": "企业获得订单及交付状态",
    "team_material_orders": "企业在途原材料",
    "team_material_inventory": "企业原材料库存",
    "team_product_inventory": "企业产成品库存",
    "team_receivables": "企业应收款批次",
    "team_loans": "企业贷款明细",
    "team_qualifications": "市场、产品和 ISO 研发认证状态",
    "team_factories": "企业厂房状态",
    "team_production_lines": "企业生产线状态",
    "test_cash_flow_events": "710W 测试方案的人工整理现金流",
}


class DataPipeline:
    def __init__(self, root: Path, output: Path) -> None:
        self.root = root.resolve()
        self.output = output.resolve()
        self.rows: dict[str, list[dict[str, Any]]] = defaultdict(list)
        self.issues: list[dict[str, Any]] = []
        self._source_by_path: dict[Path, str] = {}

    def add(self, table: str, **record: Any) -> None:
        self.rows[table].append(record)

    def relative(self, path: Path) -> str:
        resolved = path.resolve()
        try:
            return resolved.relative_to(self.root).as_posix()
        except ValueError:
            # Raw资料可以通过仓库外部软链接挂载。保留可读的相对审计路径，
            # 不把外部绝对路径写入标准数据，且不因软链接越出仓库而中断导入。
            return os.path.relpath(resolved, self.root).replace(os.sep, "/")

    def source_for(self, path: Path) -> str:
        return self._source_by_path[path.resolve()]

    def _first_existing_dir(self, *relative_paths: str) -> Path:
        candidates = [self.root / relative_path for relative_path in relative_paths]
        for candidate in candidates:
            if candidate.is_dir():
                return candidate.resolve()
        joined = ", ".join(str(path) for path in candidates)
        raise FileNotFoundError(f"未找到任何候选数据目录：{joined}")

    def _historical_dir(self) -> Path:
        return self._first_existing_dir(
            "data/original/VPD-OE_Agent_20260709/VPD-OE_Agent/data/raw/ZY",
            "data/original/一场比赛数据示例",
            "一场比赛数据示例",
        )

    def _test_data_dir(self) -> Path:
        return self._first_existing_dir(
            "data/original/测试数据",
            "测试数据",
            "比赛赛前已知数据",
        )

    def build(self) -> dict[str, Any]:
        self._build_source_manifest()
        self._inventory_workbooks()
        self._parse_rules()
        self._parse_order_catalog()
        self._parse_public_year_files()
        self._parse_team_files()
        self._parse_test_cash_flow()
        self._build_rulepack_artifacts()
        frames = self._frames()
        quality = self._quality_checks(frames)
        self._write_outputs(frames, quality)
        return quality

    def _classify_source(self, rel: str) -> tuple[str, str | None, str, str]:
        suffix = Path(rel).suffix.lower()
        source_type = {
            ".xls": "excel_xls",
            ".xlsx": "excel_xlsx",
            ".docx": "docx",
            ".pdf": "pdf",
            ".md": "markdown",
            ".txt": "text",
            ".png": "image",
            ".rar": "archive",
        }.get(suffix, "other")
        if rel.startswith("一场比赛数据示例/") or "/data/raw/ZY/" in f"/{rel}":
            return HISTORICAL_COMPETITION, "unknown", "confirmed_internal_group", source_type
        if rel.endswith("试题题面01.docx"):
            return TEST_COMPETITION, "zhejiang_8th_rules_v1", "confirmed", source_type
        if rel.endswith("9手p1p4产品2937分.xlsx") or rel.endswith("五年现金流.md"):
            return TEST_COMPETITION, "zhejiang_8th_rules_v1", "probable", source_type
        if rel.endswith("原订单规则.xls"):
            return ORDER_COMPETITION, None, "unknown", source_type
        if rel.startswith("ranking_regression_report"):
            return HISTORICAL_COMPETITION, "unknown", "probable", source_type
        return DESIGN_SCOPE, None, "reference_only", source_type

    def _build_source_manifest(self) -> None:
        origin_root = self.root / "data/original"
        design_files = [path for path in origin_root.iterdir() if path.is_file()] if origin_root.is_dir() else []
        source_dirs = (self._historical_dir(), self._test_data_dir())
        files = design_files + [path for directory in source_dirs for path in directory.iterdir() if path.is_file()]
        files = [path for path in files if not path.name.startswith(".~") and path.name != "pyproject.toml"]
        files = list({path.resolve(): path for path in files}.values())
        files.sort(key=lambda path: ("测试数据" in path.parts, self.relative(path)))
        first_by_hash: dict[str, str] = {}
        for path in files:
            rel = self.relative(path)
            digest = sha256_file(path)
            sid = source_id(path, self.root)
            self._source_by_path[path.resolve()] = sid
            competition_id, rule_version, binding_status, source_type = self._classify_source(rel)
            duplicate_of = first_by_hash.get(digest)
            first_by_hash.setdefault(digest, sid)
            self.add(
                "source_manifest",
                source_id=sid,
                relative_path=rel,
                file_name=path.name,
                extension=path.suffix.lower(),
                source_type=source_type,
                size_bytes=path.stat().st_size,
                sha256=digest,
                competition_id=competition_id,
                rule_version=rule_version,
                binding_status=binding_status,
                duplicate_of_source_id=duplicate_of,
                modified_at=datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
            )

    def _inventory_workbooks(self) -> None:
        excel_files = sorted(
            path
            for path in self._source_by_path
            if path.suffix.lower() in {".xls", ".xlsx"} and not path.name.startswith(".~")
        )
        for path in excel_files:
            try:
                if path.suffix.lower() == ".xls":
                    self._inventory_xls(path)
                else:
                    self._inventory_xlsx(path)
            except Exception as exc:  # 数据保真清单不能阻断其他文件
                self.issues.append(
                    {
                        "severity": "error",
                        "code": "workbook_parse_error",
                        "message": f"无法读取 {self.relative(path)}: {exc}",
                        "count": 1,
                    }
                )

    def _inventory_xls(self, path: Path) -> None:
        sid = self.source_for(path)
        wb = xlrd.open_workbook(str(path), formatting_info=True)
        for sheet in wb.sheets():
            self.add(
                "workbook_inventory",
                source_id=sid,
                source_path=self.relative(path),
                workbook_format="xls",
                sheet_name=sheet.name,
                row_count=sheet.nrows,
                column_count=sheet.ncols,
                merged_range_count=len(sheet.merged_cells),
            )
            for row in range(sheet.nrows):
                for col in range(sheet.ncols):
                    cell = sheet.cell(row, col)
                    if cell.ctype in {xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK} or is_blank(cell.value):
                        continue
                    value_type = {
                        xlrd.XL_CELL_TEXT: "text",
                        xlrd.XL_CELL_NUMBER: "number",
                        xlrd.XL_CELL_DATE: "date",
                        xlrd.XL_CELL_BOOLEAN: "boolean",
                        xlrd.XL_CELL_ERROR: "error",
                    }.get(cell.ctype, "unknown")
                    value = cell.value
                    if cell.ctype == xlrd.XL_CELL_DATE:
                        value = xlrd.xldate_as_datetime(value, wb.datemode).isoformat()
                    self.add(
                        "workbook_cells",
                        source_id=sid,
                        source_path=self.relative(path),
                        sheet_name=sheet.name,
                        source_cell=cell_ref(row, col),
                        source_row=row + 1,
                        source_column=col + 1,
                        value_type=value_type,
                        raw_value=clean_text(value),
                        formula=None,
                    )

    def _inventory_xlsx(self, path: Path) -> None:
        sid = self.source_for(path)
        wb = load_workbook(path, data_only=False, read_only=False)
        for sheet in wb.worksheets:
            self.add(
                "workbook_inventory",
                source_id=sid,
                source_path=self.relative(path),
                workbook_format="xlsx",
                sheet_name=sheet.title,
                row_count=sheet.max_row,
                column_count=sheet.max_column,
                merged_range_count=len(sheet.merged_cells.ranges),
            )
            for row in sheet.iter_rows():
                for cell in row:
                    if is_blank(cell.value):
                        continue
                    value = cell.value
                    formula = value if isinstance(value, str) and value.startswith("=") else None
                    if isinstance(value, (datetime, date)):
                        value = value.isoformat()
                    self.add(
                        "workbook_cells",
                        source_id=sid,
                        source_path=self.relative(path),
                        sheet_name=sheet.title,
                        source_cell=cell.coordinate,
                        source_row=cell.row,
                        source_column=cell.column,
                        value_type="formula" if formula else cell.data_type,
                        raw_value=clean_text(value),
                        formula=formula,
                    )

    def _primary_path(self, name: str) -> Path:
        path = self._test_data_dir() / name
        if not path.is_file():
            raise FileNotFoundError(f"测试数据中缺少文件：{path}")
        return path

    def _parse_rules(self) -> None:
        path = self._primary_path("试题题面01.docx")
        document = Document(path)
        sid = self.source_for(path)
        all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        initial_match = re.search(r"初始资本\s*(\d+(?:\.\d+)?)\s*W", all_text)
        management_match = re.search(r"管理费\s*(\d+(?:\.\d+)?)\s*W", all_text)
        self.add(
            "competition_rules",
            competition_id=TEST_COMPETITION,
            rule_version="zhejiang_8th_rules_v1",
            competition_name="新道杯第八届浙江省大学生企业经营沙盘模拟竞赛（本科组）",
            initial_capital_wan=float(initial_match.group(1)) if initial_match else None,
            management_fee_per_quarter_wan=float(management_match.group(1)) if management_match else None,
            source_id=sid,
            source_path=self.relative(path),
        )

        finance_table = document.tables[0]
        for row_index, row in enumerate(finance_table.rows[1:], start=2):
            values = [cell.text.strip() for cell in row.cells]
            rate = as_number(values[1])
            max_term = re.search(r"最长可贷年限\s*(\d+)\s*年", values[1])
            self.add(
                "rule_financing",
                competition_id=TEST_COMPETITION,
                rule_version="zhejiang_8th_rules_v1",
                finance_type=values[0],
                annual_rate=rate / 100 if rate is not None else None,
                max_term_years=int(max_term.group(1)) if max_term else None,
                raw_rule=values[1],
                source_id=sid,
                source_path=self.relative(path),
                source_table=1,
                source_row=row_index,
            )

        financing_terms = [
            ("long_loan", "长期贷款", 0.10, "annual", None, 4, "year", 2),
            ("short_loan", "短期贷款", 0.05, "annual", None, None, None, 3),
            ("discount_1_2q", "资金贴现", 0.08, "as_stated", 1, 2, "quarter", 4),
            ("discount_3_4q", "资金贴现", 0.11, "as_stated", 3, 4, "quarter", 4),
        ]
        for term_id, finance_type, rate, rate_basis, term_min, term_max, term_unit, source_row in financing_terms:
            self.add(
                "rule_financing_terms",
                term_id=term_id,
                competition_id=TEST_COMPETITION,
                rule_version="zhejiang_8th_rules_v1",
                finance_type=finance_type,
                rate=rate,
                rate_basis=rate_basis,
                term_min=term_min,
                term_max=term_max,
                term_unit=term_unit,
                semantic_status="explicit_rate_partial_process",
                source_id=sid,
                source_path=self.relative(path),
                source_table=1,
                source_row=source_row,
            )

        self._parse_rule_table(document, sid, path, 1, "rule_factories", [
            "factory_id", "factory_name", "purchase_price_wan", "annual_rent_wan", "sale_price_wan",
            "line_capacity", "usage_limit", "score",
        ])
        self._parse_rule_table(document, sid, path, 2, "rule_production_lines", [
            "line_type_id", "line_type_name", "investment_per_quarter_wan", "installation_quarters",
            "production_cycle_quarters", "conversion_fee_per_quarter_wan", "conversion_cycle_quarters",
            "maintenance_fee_wan", "residual_value_wan", "depreciation_fee_wan", "depreciation_years", "score",
        ])
        self._parse_rule_table(document, sid, path, 3, "rule_markets", [
            "market_id", "market_name", "annual_development_fee_wan", "development_years", "score",
        ])
        self._parse_rule_table(document, sid, path, 4, "rule_iso", [
            "iso_id", "iso_name", "annual_development_fee_wan", "development_years", "score",
        ])
        self._parse_rule_table(document, sid, path, 5, "rule_materials", [
            "material_id", "material_name", "unit_purchase_price_wan", "lead_time_quarters",
        ])
        self._parse_rule_table(document, sid, path, 6, "rule_products", [
            "product_id", "product_name", "processing_fee_wan", "development_fee_per_quarter_wan",
            "development_quarters", "direct_cost_wan", "score",
        ])
        self._parse_rule_table(document, sid, path, 7, "rule_bom", [
            "product_name", "component_product_name", "material_name", "quantity",
        ], ids_are_numeric=False)

    def _parse_rule_table(
        self,
        document: Document,
        sid: str,
        path: Path,
        table_index: int,
        output_table: str,
        columns: list[str],
        ids_are_numeric: bool = True,
    ) -> None:
        table = document.tables[table_index]
        for row_index, row in enumerate(table.rows[1:], start=2):
            raw_values = [cell.text.strip() for cell in row.cells]
            record: dict[str, Any] = {}
            for index, column in enumerate(columns):
                raw = raw_values[index] if index < len(raw_values) else None
                if column.endswith("_name"):
                    record[column] = clean_text(raw)
                elif column.endswith("_id") and not ids_are_numeric:
                    record[column] = clean_text(raw)
                else:
                    number = as_number(raw)
                    record[column] = int(number) if number is not None and number.is_integer() else number
            record.update(
                competition_id=TEST_COMPETITION,
                rule_version="zhejiang_8th_rules_v1",
                source_id=sid,
                source_path=self.relative(path),
                source_table=table_index + 1,
                source_row=row_index,
            )
            self.add(output_table, **record)

    def _parse_order_catalog(self) -> None:
        path = self._primary_path("原订单规则.xls")
        sid = self.source_for(path)
        sheet = xlrd.open_workbook(str(path)).sheet_by_name("订单一览表")
        headers = {clean_text(sheet.cell_value(0, col)): col for col in range(sheet.ncols)}
        for row in range(1, sheet.nrows):
            order_id = clean_text(sheet.cell_value(row, headers["订单号"]))
            if not order_id:
                continue
            market_code = as_int(sheet.cell_value(row, headers["市场"]))
            product_code = as_int(sheet.cell_value(row, headers["产品"]))
            iso_code = as_int(sheet.cell_value(row, headers["ISO"]))
            quantity = as_number(sheet.cell_value(row, headers["数量"]))
            amount = parse_money_wan(sheet.cell_value(row, headers["金额"]))
            self.add(
                "order_catalog",
                order_id=order_id,
                competition_id=ORDER_COMPETITION,
                rule_version=None,
                year=as_int(sheet.cell_value(row, headers["年份"])),
                market_code=market_code,
                market_id=f"M{market_code}" if market_code is not None else None,
                market_name_candidate=MARKET_NAMES.get(market_code),
                product_code=product_code,
                product_id=PRODUCT_NAMES.get(product_code),
                quantity=quantity,
                total_amount_wan=amount,
                unit_price_wan=amount / quantity if amount is not None and quantity else None,
                receivable_term_quarters=as_int(sheet.cell_value(row, headers["账期"])),
                delivery_term_quarters=as_int(sheet.cell_value(row, headers["交货期"])),
                iso_code=iso_code,
                iso_name_candidate=ISO_NAMES.get(iso_code),
                mapping_status="probable",
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )

    def _parse_public_year_files(self) -> None:
        for path in sorted(self._historical_dir().glob("[1-6].xls")):
            year = int(path.stem)
            workbook = xlrd.open_workbook(str(path))
            self._parse_public_advertising(path, workbook, year)
            self._parse_public_financials(path, workbook, year)
            self._parse_public_lines(path, workbook, year)
            self._parse_public_market_leaders(path, workbook, year)

    def _parse_public_advertising(self, path: Path, workbook: xlrd.book.Book, year: int) -> None:
        sheet_name = f"第{year}年广告投放(格式二)"
        if sheet_name not in workbook.sheet_names():
            return
        sheet = workbook.sheet_by_name(sheet_name)
        sid = self.source_for(path)
        teams = [clean_text(sheet.cell_value(0, col)) for col in range(1, sheet.ncols)]
        market: str | None = None
        for row in range(1, sheet.nrows):
            label = clean_text(sheet.cell_value(row, 0))
            if label and label.endswith("广告投放情况"):
                market = label.removesuffix("广告投放情况")
                continue
            if not label or not label.startswith("P") or market is None:
                continue
            for col, team_id in enumerate(teams, start=1):
                if not team_id:
                    continue
                amount = parse_money_wan(sheet.cell_value(row, col))
                self.add(
                    "annual_advertising",
                    competition_id=HISTORICAL_COMPETITION,
                    rule_version="unknown",
                    year=year,
                    team_id=team_id,
                    product_id=label,
                    market_name=market,
                    amount_wan=amount,
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_cell=cell_ref(row, col),
                    source_row=row + 1,
                )

    def _parse_public_financials(self, path: Path, workbook: xlrd.book.Book, year: int) -> None:
        sheet = workbook.sheet_by_name(f"第{year}年三张报表")
        sid = self.source_for(path)
        sections = [(1, 2, 13, "comprehensive_expense"), (14, 15, 26, "income_statement"), (27, 28, 50, "balance_sheet")]
        for header_row, start, stop, statement in sections:
            teams = [clean_text(sheet.cell_value(header_row, col)) for col in range(2, sheet.ncols)]
            for row in range(start, min(stop, sheet.nrows)):
                metric = clean_text(sheet.cell_value(row, 1))
                if not metric or metric in {"用户名", "类型"}:
                    continue
                for col, team_id in enumerate(teams, start=2):
                    if not team_id or is_blank(sheet.cell_value(row, col)):
                        continue
                    self.add(
                        "annual_financial_metrics",
                        competition_id=HISTORICAL_COMPETITION,
                        rule_version="unknown",
                        year=year,
                        team_id=team_id,
                        statement=statement,
                        metric_name=metric,
                        amount_wan=parse_money_wan(sheet.cell_value(row, col)),
                        source_id=sid,
                        source_path=self.relative(path),
                        source_sheet=sheet.name,
                        source_cell=cell_ref(row, col),
                        source_row=row + 1,
                    )

    def _parse_public_lines(self, path: Path, workbook: xlrd.book.Book, year: int) -> None:
        sheet = workbook.sheet_by_name("生产线信息")
        sid = self.source_for(path)
        headers = {clean_text(sheet.cell_value(1, col)): col for col in range(sheet.ncols)}
        for row in range(2, sheet.nrows):
            team_id = clean_text(sheet.cell_value(row, headers["所属用户"]))
            if not team_id:
                continue
            build_year, build_quarter = parse_period(sheet.cell_value(row, headers["建成时间"]))
            start_year, start_quarter = parse_period(sheet.cell_value(row, headers["开产时间"]))
            self.add(
                "annual_production_lines",
                record_id=stable_record_id(self.relative(path), sheet.name, row + 1),
                competition_id=HISTORICAL_COMPETITION,
                rule_version="unknown",
                snapshot_year=year,
                team_id=team_id,
                line_type=clean_text(sheet.cell_value(row, headers["名称"])),
                factory=clean_text(sheet.cell_value(row, headers["厂房"])),
                product_id=clean_text(sheet.cell_value(row, headers["产品"])),
                state=clean_text(sheet.cell_value(row, headers["状态"])),
                accumulated_depreciation_wan=parse_money_wan(sheet.cell_value(row, headers["累计折旧"])),
                production_start_year=start_year,
                production_start_quarter=start_quarter,
                conversion_time_raw=clean_text(sheet.cell_value(row, headers["转产时间"])),
                remaining_time_raw=clean_text(sheet.cell_value(row, headers["剩余时间"])),
                completion_year=build_year,
                completion_quarter=build_quarter,
                construction_start_raw=clean_text(sheet.cell_value(row, headers["开建时间"])),
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )

    def _parse_public_market_leaders(self, path: Path, workbook: xlrd.book.Book, year: int) -> None:
        sheet = workbook.sheet_by_name(f"第{year}年市场老大")
        sid = self.source_for(path)
        for row in range(1, sheet.nrows):
            market = clean_text(sheet.cell_value(row, 1))
            if not market:
                continue
            self.add(
                "annual_market_leaders",
                competition_id=HISTORICAL_COMPETITION,
                rule_version="unknown",
                year=year,
                market_name=market,
                team_id=clean_text(sheet.cell_value(row, 2)) if sheet.ncols > 2 else None,
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )

    def _parse_team_files(self) -> None:
        for path in sorted(self._historical_dir().glob("ZY[0-9][0-9].xls")):
            workbook = xlrd.open_workbook(str(path))
            team_id = path.stem
            self._parse_team_profile(path, workbook, team_id)
            self._parse_team_inventory(path, workbook, team_id)
            self._parse_team_financing(path, workbook, team_id)
            self._parse_team_qualifications(path, workbook, team_id)
            self._parse_team_assets(path, workbook, team_id)
            self._parse_team_orders(path, workbook, team_id)
            self._parse_team_cash_flows(path, workbook, team_id)
            self._parse_team_financials(path, workbook, team_id)
            self._parse_team_advertising(path, workbook, team_id)

    @staticmethod
    def _value_right_of(sheet: xlrd.sheet.Sheet, label: str) -> Any:
        for row in range(sheet.nrows):
            for col in range(sheet.ncols - 1):
                if clean_text(sheet.cell_value(row, col)) == label:
                    return sheet.cell_value(row, col + 1)
        return None

    def _parse_team_profile(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("企业信息")
        cash_sheet = workbook.sheet_by_name("现金流量表")
        sid = self.source_for(path)
        system_year, system_quarter = parse_period(self._value_right_of(sheet, "系统时间"))
        has_transactions = cash_sheet.nrows > 3 and any(
            not is_blank(cash_sheet.cell_value(row, 2)) for row in range(3, cash_sheet.nrows)
        )
        self.add(
            "teams",
            competition_id=HISTORICAL_COMPETITION,
            rule_version="unknown",
            team_id=team_id,
            company_name=clean_text(self._value_right_of(sheet, "公司名称")),
            school_name=clean_text(self._value_right_of(sheet, "所属学校")),
            company_status=clean_text(self._value_right_of(sheet, "公司状态")),
            current_cash_wan=parse_money_wan(self._value_right_of(sheet, "公司现金")),
            capital_injection_wan=parse_money_wan(self._value_right_of(sheet, "股东注资")),
            system_year=system_year,
            system_quarter=system_quarter,
            data_status="active" if has_transactions else "empty_export",
            source_id=sid,
            source_path=self.relative(path),
            source_sheet=sheet.name,
        )

    def _parse_team_inventory(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("库存信息")
        sid = self.source_for(path)
        for row in range(3, sheet.nrows):
            material = clean_text(sheet.cell_value(row, 1))
            if material:
                order_year, order_quarter = parse_period(sheet.cell_value(row, 4))
                remaining, unit, _ = parse_duration(sheet.cell_value(row, 3))
                self.add(
                    "team_material_orders",
                    record_id=stable_record_id(team_id, sheet.name, "material_order", row + 1),
                    competition_id=HISTORICAL_COMPETITION,
                    team_id=team_id,
                    material_id=material,
                    quantity=as_number(sheet.cell_value(row, 2)),
                    remaining_time=remaining,
                    remaining_time_unit=unit,
                    order_year=order_year,
                    order_quarter=order_quarter,
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_row=row + 1,
                )
            material_stock = clean_text(sheet.cell_value(row, 6)) if sheet.ncols > 7 else None
            if material_stock:
                self.add(
                    "team_material_inventory",
                    record_id=stable_record_id(team_id, sheet.name, "material_stock", row + 1),
                    competition_id=HISTORICAL_COMPETITION,
                    team_id=team_id,
                    material_id=material_stock,
                    quantity=as_number(sheet.cell_value(row, 7)),
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_row=row + 1,
                )
            product_stock = clean_text(sheet.cell_value(row, 9)) if sheet.ncols > 10 else None
            if product_stock:
                self.add(
                    "team_product_inventory",
                    record_id=stable_record_id(team_id, sheet.name, "product_stock", row + 1),
                    competition_id=HISTORICAL_COMPETITION,
                    team_id=team_id,
                    product_id=product_stock,
                    quantity=as_number(sheet.cell_value(row, 10)),
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_row=row + 1,
                )

    def _parse_team_financing(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("银行贷款")
        sid = self.source_for(path)
        for row in range(3, sheet.nrows):
            receivable_amount = parse_money_wan(sheet.cell_value(row, 2))
            if receivable_amount is not None:
                term, unit, _ = parse_duration(sheet.cell_value(row, 1))
                self.add(
                    "team_receivables",
                    record_id=stable_record_id(team_id, "receivable", row + 1),
                    competition_id=HISTORICAL_COMPETITION,
                    team_id=team_id,
                    remaining_term=term,
                    remaining_term_unit=unit,
                    amount_wan=receivable_amount,
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_row=row + 1,
                )
            for loan_type, term_col, amount_col, time_col in [
                ("long", 4, 5, 6),
                ("short", 8, 9, 10),
                ("special", None, 12, 13),
            ]:
                amount = parse_money_wan(sheet.cell_value(row, amount_col)) if amount_col < sheet.ncols else None
                if amount is None:
                    continue
                remaining, unit, _ = parse_duration(sheet.cell_value(row, term_col)) if term_col is not None else (None, None, False)
                start_year, start_quarter = parse_period(sheet.cell_value(row, time_col))
                self.add(
                    "team_loans",
                    record_id=stable_record_id(team_id, loan_type, row + 1),
                    competition_id=HISTORICAL_COMPETITION,
                    team_id=team_id,
                    loan_type=loan_type,
                    amount_wan=amount,
                    remaining_term=remaining,
                    remaining_term_unit=unit,
                    start_year=start_year,
                    start_quarter=start_quarter,
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_row=row + 1,
                )

    def _parse_team_qualifications(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("研发认证")
        sid = self.source_for(path)
        blocks = [("market", 1), ("product", 7), ("iso", 13)]
        for qualification_type, start_col in blocks:
            for row in range(3, sheet.nrows):
                name = clean_text(sheet.cell_value(row, start_col)) if start_col < sheet.ncols else None
                if not name:
                    continue
                cycle, cycle_unit, _ = parse_duration(sheet.cell_value(row, start_col + 2))
                remaining, remaining_unit, complete = parse_duration(sheet.cell_value(row, start_col + 3))
                completion_year, completion_quarter = parse_period(sheet.cell_value(row, start_col + 4))
                self.add(
                    "team_qualifications",
                    record_id=stable_record_id(team_id, qualification_type, name),
                    competition_id=HISTORICAL_COMPETITION,
                    team_id=team_id,
                    qualification_type=qualification_type,
                    qualification_name=name,
                    development_fee_wan=parse_money_wan(sheet.cell_value(row, start_col + 1)),
                    development_cycle=cycle,
                    development_cycle_unit=cycle_unit,
                    remaining_time=remaining,
                    remaining_time_unit=remaining_unit,
                    is_complete=complete,
                    completion_year=completion_year,
                    completion_quarter=completion_quarter,
                    source_id=sid,
                    source_path=self.relative(path),
                    source_sheet=sheet.name,
                    source_row=row + 1,
                )

    def _parse_team_assets(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("厂房与生产线")
        sid = self.source_for(path)
        line_title_row = next(
            (row for row in range(sheet.nrows) if clean_text(sheet.cell_value(row, 1)) == "生产线信息"),
            sheet.nrows,
        )
        for row in range(3, line_title_row):
            factory_id = as_int(sheet.cell_value(row, 1))
            if factory_id is None:
                continue
            self.add(
                "team_factories",
                record_id=stable_record_id(team_id, "factory", factory_id),
                competition_id=HISTORICAL_COMPETITION,
                team_id=team_id,
                factory_instance_id=factory_id,
                factory_name=clean_text(sheet.cell_value(row, 2)),
                state=clean_text(sheet.cell_value(row, 3)),
                capacity_raw=clean_text(sheet.cell_value(row, 4)),
                purchase_price_wan=parse_money_wan(sheet.cell_value(row, 5)),
                annual_rent_wan=parse_money_wan(sheet.cell_value(row, 6)),
                sale_price_wan=parse_money_wan(sheet.cell_value(row, 7)),
                last_rent_payment_raw=clean_text(sheet.cell_value(row, 8)),
                acquisition_time_raw=clean_text(sheet.cell_value(row, 9)),
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )
        line_start = line_title_row + 2
        for row in range(line_start, sheet.nrows):
            line_id = as_int(sheet.cell_value(row, 1))
            if line_id is None:
                continue
            completion_year, completion_quarter = parse_period(sheet.cell_value(row, 10))
            start_year, start_quarter = parse_period(sheet.cell_value(row, 7))
            self.add(
                "team_production_lines",
                record_id=stable_record_id(team_id, "line", line_id),
                competition_id=HISTORICAL_COMPETITION,
                team_id=team_id,
                line_instance_id=line_id,
                line_type=clean_text(sheet.cell_value(row, 2)),
                factory=clean_text(sheet.cell_value(row, 3)),
                product_id=clean_text(sheet.cell_value(row, 4)),
                state=clean_text(sheet.cell_value(row, 5)),
                accumulated_depreciation_wan=parse_money_wan(sheet.cell_value(row, 6)),
                production_start_year=start_year,
                production_start_quarter=start_quarter,
                conversion_time_raw=clean_text(sheet.cell_value(row, 8)),
                remaining_time_raw=clean_text(sheet.cell_value(row, 9)),
                completion_year=completion_year,
                completion_quarter=completion_quarter,
                construction_start_raw=clean_text(sheet.cell_value(row, 11)),
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )

    def _parse_team_orders(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("订单信息")
        sid = self.source_for(path)
        for row in range(3, sheet.nrows):
            order_id = clean_text(sheet.cell_value(row, 1))
            if not order_id:
                continue
            won_year, _ = parse_period(sheet.cell_value(row, 7))
            delivered_year, delivered_quarter = parse_period(sheet.cell_value(row, 11))
            delivery_term, _, _ = parse_duration(sheet.cell_value(row, 8))
            receivable_term, _, _ = parse_duration(sheet.cell_value(row, 9))
            self.add(
                "team_orders",
                record_id=stable_record_id(team_id, order_id),
                competition_id=HISTORICAL_COMPETITION,
                team_id=team_id,
                order_id=order_id,
                market_name=clean_text(sheet.cell_value(row, 2)),
                product_id=clean_text(sheet.cell_value(row, 3)),
                quantity=as_number(sheet.cell_value(row, 4)),
                total_amount_wan=parse_money_wan(sheet.cell_value(row, 5)),
                status=clean_text(sheet.cell_value(row, 6)),
                won_year=won_year,
                delivery_term_quarters=delivery_term,
                receivable_term_quarters=receivable_term,
                iso_name=clean_text(sheet.cell_value(row, 10)),
                delivered_year=delivered_year,
                delivered_quarter=delivered_quarter,
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )

    def _parse_team_cash_flows(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("现金流量表")
        sid = self.source_for(path)
        for row in range(3, sheet.nrows):
            action = clean_text(sheet.cell_value(row, 2))
            if not action:
                continue
            year, quarter = parse_period(sheet.cell_value(row, 5))
            transaction_id = as_int(sheet.cell_value(row, 1))
            self.add(
                "team_cash_flows",
                record_id=stable_record_id(team_id, transaction_id, row + 1),
                competition_id=HISTORICAL_COMPETITION,
                team_id=team_id,
                transaction_id=transaction_id,
                sequence_in_source=row - 2,
                year=year,
                quarter=quarter,
                action=action,
                amount_wan=parse_money_wan(sheet.cell_value(row, 3)),
                balance_wan=parse_money_wan(sheet.cell_value(row, 4)),
                note=clean_text(sheet.cell_value(row, 6)),
                source_id=sid,
                source_path=self.relative(path),
                source_sheet=sheet.name,
                source_row=row + 1,
            )

    def _parse_team_financials(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("三张报表")
        sid = self.source_for(path)
        sections = [(1, 2, 13, "comprehensive_expense"), (14, 15, 26, "income_statement"), (27, 28, 50, "balance_sheet")]
        for header_row, start, stop, statement in sections:
            periods = [clean_text(sheet.cell_value(header_row, col)) for col in range(2, sheet.ncols)]
            for row in range(start, min(stop, sheet.nrows)):
                metric = clean_text(sheet.cell_value(row, 1))
                if not metric or metric in {"年度", "类型"}:
                    continue
                for col, period in enumerate(periods, start=2):
                    if not period or is_blank(sheet.cell_value(row, col)):
                        continue
                    year, _ = parse_period(period)
                    self.add(
                        "team_financial_metrics",
                        competition_id=HISTORICAL_COMPETITION,
                        team_id=team_id,
                        year=year,
                        period_label=period,
                        statement=statement,
                        metric_name=metric,
                        amount_wan=parse_money_wan(sheet.cell_value(row, col)),
                        source_id=sid,
                        source_path=self.relative(path),
                        source_sheet=sheet.name,
                        source_cell=cell_ref(row, col),
                        source_row=row + 1,
                    )

    def _parse_team_advertising(self, path: Path, workbook: xlrd.book.Book, team_id: str) -> None:
        sheet = workbook.sheet_by_name("广告投放")
        sid = self.source_for(path)
        row = 0
        while row < sheet.nrows:
            title = clean_text(sheet.cell_value(row, 1)) if sheet.ncols > 1 else None
            match = re.search(r"第(\d+)年广告投放情况", title or "")
            if not match:
                row += 1
                continue
            year = int(match.group(1))
            header_row = row + 1
            markets = [clean_text(sheet.cell_value(header_row, col)) for col in range(2, sheet.ncols)]
            data_row = header_row + 1
            while data_row < sheet.nrows:
                product = clean_text(sheet.cell_value(data_row, 1))
                if not product or not product.startswith("P"):
                    break
                for col, market in enumerate(markets, start=2):
                    if not market:
                        continue
                    self.add(
                        "team_advertising",
                        competition_id=HISTORICAL_COMPETITION,
                        team_id=team_id,
                        year=year,
                        product_id=product,
                        market_name=market,
                        amount_wan=parse_money_wan(sheet.cell_value(data_row, col)),
                        source_id=sid,
                        source_path=self.relative(path),
                        source_sheet=sheet.name,
                        source_cell=cell_ref(data_row, col),
                        source_row=data_row + 1,
                    )
                data_row += 1
            row = data_row

    def _parse_test_cash_flow(self) -> None:
        path = self._primary_path("五年现金流.md")
        sid = self.source_for(path)
        pattern = re.compile(
            r"^\|\s*Y(?P<year>\d+)Q(?P<quarter>\d+)\s*\|\s*(?P<action>[^|]+?)\s*\|\s*(?P<amount>[+-]?\d+(?:\.\d+)?)W\s*\|\s*\*{0,2}(?P<balance>\d+(?:\.\d+)?)W\*{0,2}\s*\|"
        )
        sequence = 0
        for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
            match = pattern.match(line)
            if not match:
                continue
            sequence += 1
            self.add(
                "test_cash_flow_events",
                record_id=stable_record_id(self.relative(path), line_number),
                competition_id=TEST_COMPETITION,
                rule_version="zhejiang_8th_rules_v1",
                scenario_id="9line_p1_p4_2937",
                sequence=sequence,
                year=int(match.group("year")),
                quarter=int(match.group("quarter")),
                action=match.group("action").strip(),
                amount_wan=float(match.group("amount")),
                balance_wan=float(match.group("balance")),
                source_id=sid,
                source_path=self.relative(path),
                source_line=line_number,
            )

    def _build_rulepack_artifacts(self) -> None:
        rule_path = self._primary_path("试题题面01.docx")
        rule_source_id = self.source_for(rule_path)
        self.add(
            "rule_packs",
            rule_pack_id="zhejiang_8th_rules_v1",
            competition_id=TEST_COMPETITION,
            rule_version="zhejiang_8th_rules_v1",
            pack_status="partial",
            parameter_status="structured_from_explicit_tables",
            process_status="blocking_gaps_present",
            order_binding_status="unbound",
            simulation_ready=False,
            blocking_gap_count=sum(severity == "blocker" for *_, severity in RULE_GAPS),
            source_id=rule_source_id,
            source_path=self.relative(rule_path),
        )
        for gap_id, domain, title, description, severity in RULE_GAPS:
            self.add(
                "rule_gaps",
                gap_id=gap_id,
                rule_pack_id="zhejiang_8th_rules_v1",
                competition_id=TEST_COMPETITION,
                rule_version="zhejiang_8th_rules_v1",
                domain=domain,
                title=title,
                description=description,
                severity=severity,
                status="unresolved",
                evidence_status="not_explicit_in_available_rule_source",
                source_id=rule_source_id,
                source_path=self.relative(rule_path),
            )

        historical_counts: dict[str, int] = defaultdict(int)
        test_counts: dict[str, int] = defaultdict(int)
        for row in self.rows["team_cash_flows"]:
            canonical = canonicalize_action(row["action"], "historical")
            if canonical:
                historical_counts[canonical] += 1
        for row in self.rows["test_cash_flow_events"]:
            canonical = canonicalize_action(row["action"], "test")
            if canonical:
                test_counts[canonical] += 1

        for canonical_action, definition in ACTION_DEFINITIONS.items():
            control_type = ACTION_CONTROL_TYPES[canonical_action]
            self.add(
                "action_definitions",
                canonical_action=canonical_action,
                category=definition["category"],
                description=definition["description"],
                control_type=control_type,
                is_agent_candidate=canonical_action in AGENT_CANDIDATE_ACTIONS,
                historical_observation_count=historical_counts.get(canonical_action, 0),
                test_observation_count=test_counts.get(canonical_action, 0),
                semantic_status="observed_not_rule_verified",
                executable_status="blocked_until_rule_confirmed",
            )
        for source_scope, aliases in (
            ("historical", HISTORICAL_ACTION_ALIASES),
            ("test", TEST_ACTION_ALIASES),
        ):
            for raw_action, canonical_action in aliases.items():
                self.add(
                    "action_aliases",
                    source_scope=source_scope,
                    raw_action=raw_action,
                    canonical_action=canonical_action,
                    mapping_status="explicit_parser_mapping",
                )
        self.add(
            "action_aliases",
            source_scope="test_pattern",
            raw_action="按订单交货(*)",
            canonical_action="order_delivery",
            mapping_status="pattern_mapping",
        )
        self.add(
            "action_aliases",
            source_scope="test_pattern",
            raw_action="贴现*q(面值)",
            canonical_action="receivable_discount",
            mapping_status="pattern_mapping",
        )

        for row in self.rows["team_cash_flows"]:
            canonical = canonicalize_action(row["action"], "historical")
            parameters, parse_status = parse_action_parameters(row["action"], row.get("note"), "historical")
            self.add(
                "action_events",
                event_id=row["record_id"],
                event_source="historical_cash_flow",
                competition_id=row["competition_id"],
                rule_version="unknown",
                team_id=row["team_id"],
                scenario_id=None,
                sequence=row["sequence_in_source"],
                year=row["year"],
                quarter=row["quarter"],
                raw_action=row["action"],
                canonical_action=canonical,
                category=ACTION_DEFINITIONS.get(canonical or "", {}).get("category"),
                control_type=ACTION_CONTROL_TYPES.get(canonical or ""),
                is_agent_candidate=canonical in AGENT_CANDIDATE_ACTIONS,
                parameters_json=json.dumps(parameters, ensure_ascii=False, sort_keys=True),
                parameter_parse_status=parse_status,
                cash_effect_wan=row["amount_wan"],
                balance_wan=row["balance_wan"],
                note=row.get("note"),
                evidence_status="observed_history_unknown_rule",
                source_id=row["source_id"],
                source_path=row["source_path"],
                source_sheet=row["source_sheet"],
                source_row=row["source_row"],
                source_line=None,
            )
        for row in self.rows["test_cash_flow_events"]:
            canonical = canonicalize_action(row["action"], "test")
            parameters, parse_status = parse_action_parameters(row["action"], None, "test")
            self.add(
                "action_events",
                event_id=row["record_id"],
                event_source="curated_test_scenario",
                competition_id=row["competition_id"],
                rule_version=row["rule_version"],
                team_id=None,
                scenario_id=row["scenario_id"],
                sequence=row["sequence"],
                year=row["year"],
                quarter=row["quarter"],
                raw_action=row["action"],
                canonical_action=canonical,
                category=ACTION_DEFINITIONS.get(canonical or "", {}).get("category"),
                control_type=ACTION_CONTROL_TYPES.get(canonical or ""),
                is_agent_candidate=canonical in AGENT_CANDIDATE_ACTIONS,
                parameters_json=json.dumps(parameters, ensure_ascii=False, sort_keys=True),
                parameter_parse_status=parse_status,
                cash_effect_wan=row["amount_wan"],
                balance_wan=row["balance_wan"],
                note=None,
                evidence_status="curated_test_scenario_probable_binding",
                source_id=row["source_id"],
                source_path=row["source_path"],
                source_sheet=None,
                source_row=None,
                source_line=row["source_line"],
            )

    def _frames(self) -> dict[str, pd.DataFrame]:
        frames: dict[str, pd.DataFrame] = {}
        for table in TABLE_DESCRIPTIONS:
            frames[table] = pd.DataFrame(self.rows.get(table, []))
        return frames

    def _quality_checks(self, frames: dict[str, pd.DataFrame]) -> dict[str, Any]:
        issues = list(self.issues)
        manifest = frames["source_manifest"]
        duplicate_count = int(manifest["duplicate_of_source_id"].notna().sum()) if not manifest.empty else 0
        if duplicate_count:
            issues.append(
                {
                    "severity": "info",
                    "code": "duplicate_source_files",
                    "message": "检测到内容完全相同的重复源文件，标准业务表仅解析主副本。",
                    "count": duplicate_count,
                }
            )

        teams = frames["teams"]
        empty_teams = teams.loc[teams["data_status"] == "empty_export", "team_id"].tolist() if not teams.empty else []
        if empty_teams:
            issues.append(
                {
                    "severity": "warning",
                    "code": "empty_team_exports",
                    "message": "企业导出文件存在但没有经营流水。",
                    "count": len(empty_teams),
                    "details": ", ".join(empty_teams),
                }
            )

        action_events = frames["action_events"]
        unmapped_actions = int(action_events["canonical_action"].isna().sum()) if not action_events.empty else 0
        parameterized_actions = (
            int((action_events["parameter_parse_status"] == "complete").sum()) if not action_events.empty else 0
        )
        issues.append(
            {
                "severity": "error" if unmapped_actions else "info",
                "code": "action_event_normalization",
                "message": "现金流动作标准化及结构化事件映射检查。",
                "count": unmapped_actions,
                "details": f"events={len(action_events)}; fully_parameterized={parameterized_actions}",
            }
        )

        rule_gaps = frames["rule_gaps"]
        blocking_rule_gaps = (
            int((rule_gaps["severity"] == "blocker").sum()) if not rule_gaps.empty else 0
        )
        issues.append(
            {
                "severity": "warning" if blocking_rule_gaps else "info",
                "code": "rule_pack_blocking_gaps",
                "message": "RulePack 中尚未由可用题面确认、会阻止完整仿真的规则缺口。",
                "count": blocking_rule_gaps,
                "details": "simulation_ready=false" if blocking_rule_gaps else "simulation_ready=true",
            }
        )

        continuity_errors = self._cash_flow_continuity_errors(frames["team_cash_flows"])
        issues.append(
            {
                "severity": "error" if continuity_errors else "info",
                "code": "cash_flow_continuity",
                "message": "现金流水逐笔余额连续性校验。",
                "count": continuity_errors,
            }
        )

        identity_errors = self._balance_sheet_identity_errors(frames["team_financial_metrics"])
        issues.append(
            {
                "severity": "error" if identity_errors else "info",
                "code": "balance_sheet_identity",
                "message": "资产总计与负债和所有者权益总计校验。",
                "count": identity_errors,
            }
        )

        public_finance_errors, public_finance_compared = self._compare_public_team_financials(frames)
        issues.append(
            {
                "severity": "error" if public_finance_errors else "info",
                "code": "public_team_financial_consistency",
                "message": "公共年度报表与企业报表交叉校验。",
                "count": public_finance_errors,
                "details": f"compared={public_finance_compared}",
            }
        )

        ad_errors, ad_compared = self._compare_public_team_advertising(frames)
        issues.append(
            {
                "severity": "error" if ad_errors else "info",
                "code": "public_team_advertising_consistency",
                "message": "公共广告巡盘与企业广告表交叉校验。",
                "count": ad_errors,
                "details": f"compared={ad_compared}",
            }
        )

        historical_capital = self._historical_initial_capitals(frames["team_cash_flows"])
        rule_capital = frames["competition_rules"]["initial_capital_wan"].dropna().unique().tolist()
        if historical_capital and rule_capital and set(historical_capital) != set(rule_capital):
            issues.append(
                {
                    "severity": "error",
                    "code": "cross_competition_initial_capital_conflict",
                    "message": "历史企业初始注资与题面初始资本不一致，已保持批次隔离。",
                    "count": 1,
                    "details": f"historical={historical_capital}; rules={rule_capital}",
                }
            )

        record_counts = {table: int(len(frame)) for table, frame in frames.items()}
        return {
            "dataset_version": "v1",
            "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
            "source_root": str(self.root),
            "output_root": str(self.output),
            "record_counts": record_counts,
            "issues": issues,
            "summary": {
                "source_files": record_counts.get("source_manifest", 0),
                "workbook_sheets": record_counts.get("workbook_inventory", 0),
                "nonempty_workbook_cells": record_counts.get("workbook_cells", 0),
                "active_teams": int((teams["data_status"] == "active").sum()) if not teams.empty else 0,
                "empty_team_exports": len(empty_teams),
                "orders": record_counts.get("order_catalog", 0),
                "cash_flow_events": record_counts.get("team_cash_flows", 0),
                "structured_action_events": record_counts.get("action_events", 0),
                "blocking_rule_gaps": blocking_rule_gaps,
                "quality_error_count": sum(
                    int(issue.get("count", 0)) for issue in issues if issue.get("severity") == "error"
                ),
            },
        }

    @staticmethod
    def _cash_flow_continuity_errors(frame: pd.DataFrame) -> int:
        if frame.empty:
            return 0
        errors = 0
        for _, group in frame.sort_values(["team_id", "sequence_in_source"]).groupby("team_id", sort=False):
            previous_balance: float | None = None
            for row in group.itertuples(index=False):
                amount = row.amount_wan
                balance = row.balance_wan
                if amount is None or balance is None or pd.isna(amount) or pd.isna(balance):
                    continue
                expected = float(amount) if previous_balance is None else previous_balance + float(amount)
                if abs(expected - float(balance)) > 1e-6:
                    errors += 1
                previous_balance = float(balance)
        return errors

    @staticmethod
    def _balance_sheet_identity_errors(frame: pd.DataFrame) -> int:
        if frame.empty:
            return 0
        balance = frame[frame["statement"] == "balance_sheet"]
        pivot = balance.pivot_table(
            index=["team_id", "year"], columns="metric_name", values="amount_wan", aggfunc="first"
        )
        required = {"资产总计", "负债和所有者权益总计"}
        if not required.issubset(pivot.columns):
            return 0
        comparable = pivot.dropna(subset=list(required))
        return int(
            ((comparable["资产总计"] - comparable["负债和所有者权益总计"]).abs() > 1e-6).sum()
        )

    @staticmethod
    def _compare_public_team_financials(frames: dict[str, pd.DataFrame]) -> tuple[int, int]:
        public = frames["annual_financial_metrics"]
        team = frames["team_financial_metrics"]
        if public.empty or team.empty:
            return 0, 0
        keys = ["team_id", "year", "statement", "metric_name"]
        merged = public[keys + ["amount_wan"]].merge(
            team[keys + ["amount_wan"]], on=keys, suffixes=("_public", "_team")
        )
        diff = (merged["amount_wan_public"] - merged["amount_wan_team"]).abs()
        return int((diff > 1e-6).sum()), int(len(merged))

    @staticmethod
    def _compare_public_team_advertising(frames: dict[str, pd.DataFrame]) -> tuple[int, int]:
        public = frames["annual_advertising"]
        team = frames["team_advertising"]
        if public.empty or team.empty:
            return 0, 0
        keys = ["team_id", "year", "product_id", "market_name"]
        merged = public[keys + ["amount_wan"]].merge(
            team[keys + ["amount_wan"]], on=keys, suffixes=("_public", "_team")
        )
        diff = (merged["amount_wan_public"] - merged["amount_wan_team"]).abs()
        return int((diff > 1e-6).sum()), int(len(merged))

    @staticmethod
    def _historical_initial_capitals(frame: pd.DataFrame) -> list[float]:
        if frame.empty:
            return []
        values = frame.loc[frame["action"] == "Pay_Capital", "amount_wan"].dropna().unique().tolist()
        return sorted(float(value) for value in values)

    def _write_outputs(self, frames: dict[str, pd.DataFrame], quality: dict[str, Any]) -> None:
        csv_dir = self.output / "csv"
        metadata_dir = self.output / "metadata"
        report_dir = self.output / "reports"
        for directory in (csv_dir, metadata_dir, report_dir):
            directory.mkdir(parents=True, exist_ok=True)

        for table, frame in frames.items():
            frame.to_csv(csv_dir / f"{table}.csv", index=False, encoding="utf-8-sig")

        sqlite_tmp = self.output / "goai.tmp.sqlite"
        sqlite_path = self.output / "goai.sqlite"
        if sqlite_tmp.exists():
            sqlite_tmp.unlink()
        with sqlite3.connect(sqlite_tmp) as connection:
            for table, frame in frames.items():
                frame.to_sql(table, connection, if_exists="replace", index=False)
                columns = set(frame.columns)
                for column in ("source_id", "competition_id", "team_id", "year", "order_id"):
                    if column in columns:
                        connection.execute(
                            f'CREATE INDEX IF NOT EXISTS "idx_{table}_{column}" ON "{table}" ("{column}")'
                        )
        os.replace(sqlite_tmp, sqlite_path)

        dictionary_rows: list[dict[str, Any]] = []
        schemas: dict[str, Any] = {}
        for table, frame in frames.items():
            schemas[table] = {
                "description": TABLE_DESCRIPTIONS.get(table),
                "record_count": int(len(frame)),
                "columns": [str(column) for column in frame.columns],
            }
            for column in frame.columns:
                dictionary_rows.append(
                    {
                        "table_name": table,
                        "table_description": TABLE_DESCRIPTIONS.get(table),
                        "column_name": column,
                        "dtype": str(frame[column].dtype),
                        "nullable": bool(frame[column].isna().any()),
                        "unit": self._column_unit(str(column)),
                        "description": self._column_description(str(column)),
                    }
                )
        pd.DataFrame(dictionary_rows).to_csv(
            metadata_dir / "data_dictionary.csv", index=False, encoding="utf-8-sig"
        )
        (metadata_dir / "schemas.json").write_text(
            json.dumps(schemas, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        rulepack_tables = [
            "rule_packs",
            "competition_rules",
            "rule_financing_terms",
            "rule_factories",
            "rule_production_lines",
            "rule_markets",
            "rule_iso",
            "rule_materials",
            "rule_products",
            "rule_bom",
            "rule_gaps",
            "action_definitions",
            "action_aliases",
        ]
        rulepack_payload = {
            "format_version": "rulepack_v0.1",
            "rule_pack_id": "zhejiang_8th_rules_v1",
            "generated_at": quality["generated_at"],
            "tables": {
                table: json.loads(frames[table].to_json(orient="records", force_ascii=False))
                for table in rulepack_tables
            },
        }
        (metadata_dir / "rulepack_zhejiang_8th_rules_v1.json").write_text(
            json.dumps(rulepack_payload, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        (report_dir / "run_summary.json").write_text(
            json.dumps(quality, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        (report_dir / "data_quality.md").write_text(self._quality_markdown(quality), encoding="utf-8")

    @staticmethod
    def _column_unit(column: str) -> str | None:
        if column.endswith("_wan"):
            return "万元"
        if column.endswith("_rate"):
            return "比例"
        if column.endswith("_quarter") or column.endswith("_quarters"):
            return "季度"
        if column.endswith("_year") or column == "year":
            return "年"
        if column in {"quantity", "line_capacity", "usage_limit", "score"}:
            return "数值"
        return None

    @staticmethod
    def _column_description(column: str) -> str:
        descriptions = {
            "source_id": "源文件稳定标识",
            "source_path": "相对项目根目录的源文件路径",
            "source_sheet": "源工作表名称",
            "source_row": "源文件一基行号",
            "source_cell": "源 Excel 单元格坐标",
            "competition_id": "隔离后的比赛或资料批次标识",
            "rule_version": "规则版本；unknown 表示尚未确认",
            "record_id": "由来源和业务键生成的稳定记录标识",
            "team_id": "参赛企业标识",
            "year": "比赛年度；0 表示初始元年",
            "quarter": "比赛季度",
        }
        return descriptions.get(column, column.replace("_", " "))

    @staticmethod
    def _quality_markdown(quality: dict[str, Any]) -> str:
        lines = [
            "# GoAI 数据质量报告",
            "",
            f"生成时间：{quality['generated_at']}",
            "",
            "## 摘要",
            "",
            "| 项目 | 数量 |",
            "| --- | ---: |",
        ]
        for key, value in quality["summary"].items():
            lines.append(f"| `{key}` | {value} |")
        lines.extend(["", "## 质量检查", "", "| 严重度 | 代码 | 数量 | 说明 | 细节 |", "| --- | --- | ---: | --- | --- |"]) 
        for issue in quality["issues"]:
            details = str(issue.get("details", "")).replace("|", "\\|")
            message = str(issue.get("message", "")).replace("|", "\\|")
            lines.append(
                f"| {issue.get('severity')} | `{issue.get('code')}` | {issue.get('count', 0)} | {message} | {details} |"
            )
        lines.extend(["", "## 表记录数", "", "| 表 | 记录数 |", "| --- | ---: |"])
        for table, count in quality["record_counts"].items():
            lines.append(f"| `{table}` | {count} |")
        lines.extend(
            [
                "",
                "## 使用限制",
                "",
                "- `historical_600_unknown_rule`、`zhejiang_8th_710` 与 `order_catalog_unbound` 保持隔离，未确认前不得跨批次连接训练或重放。",
                "- `workbook_cells` 是保真兜底表；业务分析优先使用对应的标准化长表。",
                "- 金额统一为万元；保留 `raw_value` 或来源坐标以支持复核。",
                "",
            ]
        )
        return "\n".join(lines)
